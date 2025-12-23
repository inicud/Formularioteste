import json
import time
import os
import re
from pathlib import Path
from docx import Document
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

# ======================
# CONFIGURAÇÕES
# ======================
BASE_DIR = Path("/Users/moschiat/Desktop/Lucas APP")
MODELO_DOCX = BASE_DIR / "PROCURAÇÃO - MODELO.docx"
PASTA_SAIDA = BASE_DIR / "PROCURAÇÃO_PREENCHIDA"

PASTA_SAIDA.mkdir(exist_ok=True)

JSON_PATTERN = re.compile(r"dados_formulario(\s\(\d+\))?\.json")

# ======================
# FUNÇÕES
# ======================
def carregar_json(caminho: Path) -> dict:
    try:
        with caminho.open(encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print(f"❌ Erro ao ler JSON {caminho.name}: {e}")
        return {}


def substituir_placeholders(texto: str, dados: dict) -> str:
    for chave, valor in dados.items():
        texto = texto.replace(f"<<{chave}>>", str(valor))
    return texto


def preencher_docx(json_path: Path):
    if not MODELO_DOCX.exists():
        print("❌ Modelo DOCX não encontrado.")
        return

    dados = carregar_json(json_path)
    if not dados:
        return

    try:
        doc = Document(MODELO_DOCX)
    except Exception as e:
        print(f"❌ Erro ao abrir DOCX: {e}")
        return

    # Parágrafos
    for p in doc.paragraphs:
        if "<<" in p.text:
            p.text = substituir_placeholders(p.text, dados)

    # Tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "<<" in cell.text:
                    cell.text = substituir_placeholders(cell.text, dados)

    nome_saida = json_path.stem + ".docx"
    caminho_saida = PASTA_SAIDA / nome_saida

    doc.save(caminho_saida)
    print(f"✅ Documento gerado: {caminho_saida.name}")


# ======================
# WATCHER
# ======================
class MonitorArquivos(FileSystemEventHandler):
    def on_created(self, event):
        if event.is_directory:
            return

        caminho = Path(event.src_path)

        if JSON_PATTERN.fullmatch(caminho.name):
            time.sleep(0.7)  # garante escrita completa
            preencher_docx(caminho)


# ======================
# MAIN
# ======================
if __name__ == "__main__":
    print("👀 Monitorando pasta:", BASE_DIR)

    observer = Observer()
    observer.schedule(MonitorArquivos(), str(BASE_DIR), recursive=False)
    observer.start()

    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        print("\n🛑 Monitoramento encerrado.")
        observer.stop()

    observer.join()
